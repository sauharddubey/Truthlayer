"""RAG knowledge system (FR-RAG-001..004).

Handles document parsing (PDF/DOCX/TXT), chunking, embedding, and tenant-scoped
semantic retrieval over pgvector. All retrieval is organization-scoped so business
knowledge never leaks across tenants (NFR-SEC-002).
"""

from __future__ import annotations

import io
import logging
from typing import List, Optional

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.llm import embed_one, embed_texts
from app.models import BusinessDocument, DocumentChunk

logger = logging.getLogger("truthlayer.rag")

CHUNK_SIZE = 900
CHUNK_OVERLAP = 150

PRODUCT_DETAILS = "product_details"
MARKETING_POLICY = "marketing_policy"
ALLOWED_DOCUMENT_TYPES = {PRODUCT_DETAILS, MARKETING_POLICY}
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


# ── Parsing ──────────────────────────────────────────────────────────────────


def parse_document(filename: str, raw: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _parse_pdf(raw)
    if name.endswith(".docx"):
        return _parse_docx(raw)
    # txt / md / fallback
    return raw.decode("utf-8", errors="ignore")


def _parse_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:  # pragma: no cover
        logger.exception("PDF parse failed: %s", exc)
        return ""


def _parse_docx(raw: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as exc:  # pragma: no cover
        logger.exception("DOCX parse failed: %s", exc)
        return ""


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    text = " ".join(text.split())
    if not text:
        return []
    chunks, start = [], 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks


# ── Ingestion ────────────────────────────────────────────────────────────────


def ingest_document(
    db: Session,
    *,
    organization_id: str,
    filename: str,
    document_type: str,
    raw: bytes,
    product_id: Optional[str] = None,
) -> BusinessDocument:
    if document_type not in ALLOWED_DOCUMENT_TYPES:
        raise ValueError(f"document_type must be one of {sorted(ALLOWED_DOCUMENT_TYPES)}")
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file type. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}")

    text = parse_document(filename, raw)
    chunks = chunk_text(text)

    doc = BusinessDocument(
        organization_id=organization_id,
        product_id=product_id,
        document_type=document_type,
        filename=filename,
        status="indexing",
    )
    db.add(doc)
    db.flush()

    if chunks:
        vectors = embed_texts(chunks)
        for i, (content, vec) in enumerate(zip(chunks, vectors)):
            db.add(
                DocumentChunk(
                    document_id=doc.id,
                    organization_id=organization_id,
                    product_id=product_id,
                    content=content,
                    chunk_index=i,
                    embedding=vec,
                    extra_metadata={"document_type": document_type},
                )
            )
        doc.status = "indexed"
    else:
        doc.status = "failed"
        logger.warning("No text extracted from %s (%s)", filename, document_type)
    db.commit()
    db.refresh(doc)
    return doc


# ── Retrieval ────────────────────────────────────────────────────────────────


def retrieve(
    db: Session,
    *,
    organization_id: Optional[str],
    query: str,
    k: int = 5,
    product_id: Optional[str] = None,
    document_types: Optional[List[str]] = None,
) -> List[dict]:
    """Semantic search over an organization's (optionally product's) knowledge."""
    if not organization_id:
        return []
    query_vec = embed_one(query)
    stmt = (
        select(DocumentChunk)
        .join(BusinessDocument, DocumentChunk.document_id == BusinessDocument.id)
        .where(
            DocumentChunk.organization_id == organization_id,
            BusinessDocument.status == "indexed",
        )
    )
    if product_id:
        stmt = stmt.where(DocumentChunk.product_id == product_id)
    if document_types:
        stmt = stmt.where(BusinessDocument.document_type.in_(document_types))
    stmt = stmt.order_by(DocumentChunk.embedding.cosine_distance(query_vec)).limit(k)
    rows = db.execute(stmt).scalars().all()
    return [
        {
            "content": r.content,
            "document_id": r.document_id,
            "chunk_index": r.chunk_index,
            "document_type": (r.extra_metadata or {}).get("document_type"),
        }
        for r in rows
    ]
